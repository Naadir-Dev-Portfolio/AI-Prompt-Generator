#!/usr/bin/env python3
"""
AI Prompt Toolkit
– Generates a “master prompt” in Markdown
– Parses pasted Markdown and appends it, faithfully converting bullets and numbers, to a .docx
"""

import sys
import os
import re
from pathlib import Path
from zipfile import BadZipFile

from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QTabWidget,
    QVBoxLayout, QHBoxLayout, QLabel, QComboBox,
    QSpinBox, QPushButton, QTextEdit, QCheckBox,
    QFileDialog, QMessageBox
)
from PyQt6.QtGui import QGuiApplication, QPalette, QColor
from PyQt6.QtCore import Qt

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ------------------------------------------------------------------ #
# Example titles per chapter
# ------------------------------------------------------------------ #
EXAMPLE_TITLES = {
    "Quiz & Trivia Challenges": [
        "’80s Classroom Pop-Quiz",
        "C-3PO’s Galactic Trivia Circuit",
        "Saturday-Morning Cartoon Quiz"
    ],
    "Immersive Role-Play Realms": [
        "Noir Detective Interview Roleplay",
        "Medieval Council Diplomacy Simulation",
        "Cyberpunk Undercover Negotiation"
    ],
    # … add the rest of your chapters here …
}

# ------------------------------------------------------------------ #
# Master-prompt template (Markdown with UX tags)
# ------------------------------------------------------------------ #
MASTER_TEMPLATE = r"""# Master Prompt for Chapter: {chapter_name}

You are an expert creative-writing engine hired to supply **system prompts** that instruct a host LLM to run extremely immersive experiences.  
For the chapter **{chapter_name}**, create **{count}** unique system prompts, each 200–300 words.

**Output in pure Markdown** with this structure:

UX TITLE: <Short Title>  
UX DESCRIPTION: <One-sentence description>  

Then your instructions, using Markdown:
- Bullet lists with `- `  
- Numbered lists with `1.`, `2.`, etc.  

Separate each prompt with a line containing exactly `---`.  
**Do not** include code fences or extra commentary—only the prompts themselves.

Begin now; return only the Markdown prompts and delimiters.
"""

# ------------------------------------------------------------------ #
# Regex for bolding tokens with ≥2 uppercase letters
# ------------------------------------------------------------------ #
ACRONYM_RE = re.compile(r'.*[A-Z].*[A-Z].*')

def _add_paragraph_with_acronyms(p, text):
    """
    Split text into runs and bold tokens matching ACRONYM_RE.
    """
    for token in re.split(r'(\W+)', text):
        run = p.add_run(token)
        if ACRONYM_RE.fullmatch(token):
            run.bold = True

# ------------------------------------------------------------------ #
# Parse pasted Markdown into (title, description, body) tuples
# ------------------------------------------------------------------ #
def parse_prompts(raw_md: str):
    """
    Splits on lines of '---' and within each block finds:
      UX TITLE: ...
      UX DESCRIPTION: ...
      rest is body (markdown bullets & numbers)
    """
    blocks = re.split(r'(?m)^\s*---\s*$', raw_md.strip())
    prompts = []
    for blk in blocks:
        blk = blk.strip()
        if not blk:
            continue
        lines = blk.splitlines()
        title = desc = ""
        i = 0
        # extract UX TITLE and UX DESCRIPTION
        for i, ln in enumerate(lines):
            if ln.startswith("UX TITLE:"):
                title = ln.split("UX TITLE:",1)[1].strip()
            elif ln.startswith("UX DESCRIPTION:"):
                desc = ln.split("UX DESCRIPTION:",1)[1].strip()
                i += 1
                break
        body = "\n".join(lines[i:]).strip()
        prompts.append((title, desc, body))
    return prompts

# ------------------------------------------------------------------ #
# Append prompts to a .docx with faithful bullet/number conversion
# ------------------------------------------------------------------ #
def append_to_docx(path: Path, prompts):
    """
    Each prompt becomes a new page:
      • Title, centered and styled
      • Description, italic
      • "Copy entire prompt below:" line, centered
      • Body: only lines starting with "- " become bullets;
              lines starting with "1. " etc. become literal numbered items;
              all others become normal paragraphs.
    """
    try:
        doc = Document(path)
    except BadZipFile:
        if QMessageBox.question(
            None, "Invalid File",
            f"'{path.name}' is not a valid .docx.\nOverwrite with blank?"
        ) != QMessageBox.StandardButton.Yes:
            return False
        doc = Document()
        doc.save(path)
        doc = Document(path)

    for title, desc, body in prompts:
        # page break per prompt
        doc.add_page_break()

        # Title
        p_t = doc.add_paragraph()
        p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_t = p_t.add_run(title)
        run_t.font.name = 'Comic Sans MS'
        run_t.font.size = Pt(24)
        run_t.bold = True
        run_t.font.color.rgb = RGBColor(0x00,0x80,0x40)
        doc.add_paragraph(); doc.add_paragraph()

        # Description
        p_d = doc.add_paragraph()
        run_d = p_d.add_run(desc)
        run_d.font.italic = True
        run_d.font.size = Pt(14)
        doc.add_paragraph(); doc.add_paragraph()

        # Copy instruction
        p_c = doc.add_paragraph()
        p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_c = p_c.add_run("Copy entire prompt below:")
        run_c.bold = True
        run_c.font.size = Pt(11)
        doc.add_paragraph()

        # Body: faithful conversion
        for ln in body.splitlines():
            s = ln.lstrip()

            # numbered list item?
            m_num = re.match(r'^(\d+)\.\s+(.*)', s)
            if m_num:
                num, txt = m_num.groups()
                p_n = doc.add_paragraph(style='Normal')
                _add_paragraph_with_acronyms(p_n, f"{num}. {txt}")
                continue

            # bullet item?
            m_bul = re.match(r'^-\s+(.*)', s)
            if m_bul:
                txt = m_bul.group(1)
                p_b = doc.add_paragraph(style='List Bullet')
                p_b.paragraph_format.left_indent = Inches(0.25)
                _add_paragraph_with_acronyms(p_b, txt)
                continue

            # everything else
            p = doc.add_paragraph(style='Normal')
            _add_paragraph_with_acronyms(p, s)

        # extra spacing after each prompt
        doc.add_paragraph(); doc.add_paragraph()

    try:
        doc.save(path)
    except PermissionError:
        QMessageBox.critical(
            None, "Save Error",
            f"Cannot write to '{path.name}'. Close it and retry."
        )
        return False

    return True

# ------------------------------------------------------------------ #
# Main GUI
# ------------------------------------------------------------------ #
class UnifiedApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("AI Prompt Toolkit")
        tabs = QTabWidget()
        tabs.addTab(self._master_tab(), "Master Prompt")
        tabs.addTab(self._append_tab(), "Append to Word")
        self.setCentralWidget(tabs)
        self.resize(700, 500)

    def _master_tab(self):
        w = QWidget()
        layout = QVBoxLayout(w)

        # Chapter selector
        row = QHBoxLayout()
        row.addWidget(QLabel("Chapter:"))
        self.chapter_cb = QComboBox()
        self.chapter_cb.addItems(EXAMPLE_TITLES.keys())
        self.chapter_cb.setEditable(True)
        row.addWidget(self.chapter_cb)
        row.addStretch()
        layout.addLayout(row)

        # Prompt count
        row2 = QHBoxLayout()
        row2.addWidget(QLabel("Prompts:"))
        self.count_sb = QSpinBox()
        self.count_sb.setRange(1,10)
        self.count_sb.setValue(5)
        row2.addWidget(self.count_sb)
        row2.addStretch()
        layout.addLayout(row2)

        # Copy button
        btn = QPushButton("Copy Master Prompt")
        btn.clicked.connect(self._copy_master)
        layout.addWidget(btn, alignment=Qt.AlignmentFlag.AlignLeft)
        layout.addStretch()
        return w

    def _append_tab(self):
        w = QWidget()
        v = QVBoxLayout(w)

        # File selector
        fl = QHBoxLayout()
        self.file_label = QLabel("No document selected")
        select_btn = QPushButton("Select .docx…")
        select_btn.clicked.connect(self._select_doc)
        fl.addWidget(select_btn)
        fl.addWidget(self.file_label)
        fl.addStretch()
        v.addLayout(fl)

        # Open after save?
        self.open_cb = QCheckBox("Open after save")
        self.open_cb.setChecked(False)
        v.addWidget(self.open_cb)

        # Paste area
        v.addWidget(QLabel("Paste Markdown prompts here:"))
        self.text_edit = QTextEdit()
        v.addWidget(self.text_edit)

        # Insert button
        ins_btn = QPushButton("Insert Prompts")
        ins_btn.clicked.connect(self._insert_prompts)
        v.addWidget(ins_btn, alignment=Qt.AlignmentFlag.AlignLeft)
        v.addStretch()
        return w

    def _copy_master(self):
        chap = self.chapter_cb.currentText().strip()
        cnt = self.count_sb.value()
        text = MASTER_TEMPLATE.format(chapter_name=chap, count=cnt)
        QGuiApplication.clipboard().setText(text)
        QMessageBox.information(self, "Copied",
            "Master prompt (Markdown) copied to clipboard."
        )

    def _select_doc(self):
        fn, _ = QFileDialog.getOpenFileName(
            self, "Select Word Document", "", "Word Documents (*.docx)"
        )
        if fn:
            self.doc_path = Path(fn)
            self.file_label.setText(self.doc_path.name)

    def _insert_prompts(self):
        if not hasattr(self, 'doc_path'):
            QMessageBox.warning(self, "No File", "Please select a .docx first.")
            return
        raw = self.text_edit.toPlainText().strip()
        if not raw:
            QMessageBox.warning(self, "Empty", "Nothing was pasted.")
            return

        prompts = parse_prompts(raw)
        if not prompts:
            QMessageBox.warning(self, "Parse Error",
                "No valid prompts found. Make sure you’ve used UX TITLE:/UX DESCRIPTION: and ‘---’ separators.")
            return

        if append_to_docx(self.doc_path, prompts):
            if self.open_cb.isChecked():
                try: os.startfile(self.doc_path)
                except: pass
            self.text_edit.clear()
            QMessageBox.information(self, "Done", "Prompts appended successfully.")

# ------------------------------------------------------------------ #
# Optional dark theme
# ------------------------------------------------------------------ #
def apply_dark_theme(app):
    app.setStyle("Fusion")
    p = QPalette()
    p.setColor(QPalette.ColorRole.Window,       QColor(45,45,45))
    p.setColor(QPalette.ColorRole.WindowText,   QColor(220,220,220))
    p.setColor(QPalette.ColorRole.Base,         QColor(60,60,60))
    p.setColor(QPalette.ColorRole.AlternateBase,QColor(45,45,45))
    p.setColor(QPalette.ColorRole.ToolTipBase,  QColor(220,220,220))
    p.setColor(QPalette.ColorRole.ToolTipText,  QColor(220,220,220))
    p.setColor(QPalette.ColorRole.Text,         QColor(220,220,220))
    p.setColor(QPalette.ColorRole.Button,       QColor(60,60,60))
    p.setColor(QPalette.ColorRole.ButtonText,   QColor(220,220,220))
    p.setColor(QPalette.ColorRole.Highlight,    QColor(0,120,215))
    p.setColor(QPalette.ColorRole.HighlightedText, QColor(255,255,255))
    app.setPalette(p)

# ------------------------------------------------------------------ #
# Entry point
# ------------------------------------------------------------------ #
def main():
    app = QApplication(sys.argv)
    apply_dark_theme(app)
    win = UnifiedApp()
    win.show()
    sys.exit(app.exec())

if __name__ == "__main__":
    main()
