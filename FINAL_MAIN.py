#!/usr/bin/env python3
"""
AI Prompt Toolkit
– Generates a “master prompt” in Markdown
– Parses pasted Markdown and appends it, beautifully formatted, to a .docx
"""

import sys
import os
import re
from pathlib import Path
from zipfile import BadZipFile

from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QTabWidget,
    QVBoxLayout, QHBoxLayout, QLabel, QComboBox,
    QSpinBox, QPushButton, QTextEdit, QCheckBox,
    QFileDialog, QMessageBox
)
from PyQt6.QtGui import QGuiApplication, QPalette, QColor
from PyQt6.QtCore import Qt

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ------------------------------------------------------------------ #
# Example titles per chapter (unchanged)
# ------------------------------------------------------------------ #
EXAMPLE_TITLES = {
    "Quiz & Trivia Challenges": [
        "’80s Classroom Pop‑Quiz",
        "C‑3PO’s Galactic Trivia Circuit",
        "Saturday‑Morning Cartoon Quiz"
    ],
    "Immersive Role-Play Realms": [
        "Noir Detective Interview Roleplay",
        "Medieval Council Diplomacy Simulation",
        "Cyberpunk Undercover Negotiation"
    ],
    "Choose-Your-Path Adventures": [
        "Haunted Mansion Survival Branching Tale",
        "Solar Explorer’s Moral Crossroads",
        "Time‑Loop Detective Narrative"
    ],
    "Mystery & Logic Puzzles": [
        "Locked‑Room Alibi Reconstruction",
        "Ciphered Journal of a Mad Scientist",
        "Quantum Paradox Puzzle Chamber"
    ],
    "Storycraft Forge": [
        "Epic Myth Weaving Workshop",
        "Post‑Apocalypse Colony Chronicle",
        "Shakespearean Tragedy Remix"
    ],
    "What-If Simulators": [
        "If Dinosaurs Roamed Today Scenario",
        "Alternate History Space Race",
        "Steampunk Industrial Revolution Reimagined"
    ],
    "Language Playgrounds": [
        "Shakespearean Tongue Twister Challenge",
        "Emoji Pictionary Showdown",
        "Cryptic Crossword Duel"
    ],
    "Creative Brain Boosters": [
        "Surrealist Image Association Game",
        "Random Word Story Sprint",
        "One‑Sentence Microfiction Relay"
    ],
    "Mythic & Fantasy Escapes": [
        "Dragon‑Rider’s Riddle Tournament",
        "Faerie Market Negotiation Roleplay",
        "Wizard’s Duel of Arcane Trivia"
    ],
    "Social & Party Games": [
        "Speedy Movie Quote Face‑Off",
        "Musical Mashup Guessing Game",
        "Trivia Charades Relay"
    ],
}
# ------------------------------------------------------------------ #
# Master-prompt template (outputs pure Markdown with UX tags)
# ------------------------------------------------------------------ #
MASTER_TEMPLATE = r"""# Master Prompt for Chapter: {chapter_name}

You are an expert creative‑writing engine hired to supply **system prompts** that instruct a host LLM to run **extremely immersive experiences** other wordly experiences, hypoethetical scenarios, fun and interactive sessions, etc.   
For the chapter **{chapter_name}**, create **{count}** unique system prompts, each 200–300 words.

CORE DIRECTIVES (apply to every prompt)
- **Address the LLM, not the Player.** Speak in imperatives: “Adopt a pirate persona…”, “When the Player chooses an island, present …”. Never say “You (the Player)”.
- **Distinct VOCAL STYLES:** Every prompt assigns the LLM a fresh persona (pirate, glitchy AI, ’90s host, noir detective, cartoon character, super hero, Historical Figure etc.).
- **Accents:** Vary different accents instructing the LLM to use different accents or speech styles depending on the scenario.
- **Interactive MECHANICS:** Detail how the AI should randomise questions, track score/time, offer hints, branch outcomes, escalate stakes, remember choices, etc.
- **No Boilerplate:** Vary heading names, structure, length, tone—each prompt must feel handcrafted.
- **Inclusivity & PG‑13.**
- **ENSURE THE PROMPT INCLUDES INSTRUCTIONS FOR LMM TO **NOT BREAK CHARACTER OR STRAY FROM THE EXPERIENCE** REGARDLESS OF WHAT THE USER SAYS**
- **Create very vibrant and creative scenarios from ficitious to**

**OUTPUT IN PURE MARKDOWN** with this structure:

- Prefix each prompt with `UX TITLE: <Short Title>`
- On the next line, `UX DESCRIPTION: <One-sentence description>`
- Then include your instructions, using Markdown:
  - Bullet lists with `- ` (hyphen + space)
  - Numbered lists with `1.`, `2.`, etc.
- Separate each prompt with a line containing exactly `---`
- **Do not** include code fences or extra commentary—only the prompts themselves.
- **Do not** add underscores in the UX DESCRIPTION OR UX TITLE TEXT E.G. UX_DESCRIPTION - Take extra precaution not to do this.

SECTIONS
## Section Ideas & Guidance
*(Do **not** output this section. It's for inspiration only.)*

Consider these ideas for structuring each prompt. You are **strongly encouraged** to:
- **Mix and match** elements from the following list.
- **Rename** these sections to better fit the unique context of each prompt.
- **Invent entirely new sections** that are relevant and enhance the immersive experience.

**Potential Section Elements:**

- AI Identity & Mannerisms: Define the LLM’s persona, voice, and any unique behaviors.
- Initial Scenario Setup: Describe the starting situation and any initial choices for the Player.
- Gameplay Mechanics: Detail how the core interaction (e.g., questioning, challenges) will function.
- Progression & Adaptation: Explain how the experience evolves based on Player choices or performance.
- Feedback & Consequences: Outline how the AI responds to the Player, including rewards and penalties.
- Objectives & Win/Lose States: Clearly define how the Player can succeed or fail in the experience.
- Environmental Details: Describe the setting and atmosphere of the scenario.
- Special Features & Twists: Include any unique mechanics or unexpected elements.
- Guiding the Player: Provide instructions or tips on how the Player should engage.

**Instruction:** For each of the 5 prompts you generate:
1. **Do not** include this "Section Ideas & Guidance" section in your output.
2. Select a combination of section headers that best suits the immersive experience. You are **required** to rename at least two of the suggested elements or create at least one entirely new section header for each prompt.
3. Ensure each chosen section header is bolded in your output.
4. Follow each section header with the corresponding content for that part of the prompt.
5. Omit any section headers that are not used in a particular prompt.

ADDITIONAL INSTRUCTIONS
- EACH PROMPT SHOULD HAVE A RELEVANT TITLE WITH SIMPLE NAMING ALLOWING THE USER USER TO EASILY UNDERSTAND WHAT TO EXPECT BASED ON JUST THE TITLE OF EACH PROMPT.
- EACH PROMPT SHOULD BE PROPERLY FORMATTED WITH MARKDOWN AND BULLET POINTS INCLUDING "**" TO EMPHASIZE KEY ELEMENTS IN THE PROMPT BUT AVOID CODE BLOCKS.
- DO NOT USE EMOJIS.

STYLE NOTES
- Active voice, present tense, vivid sensory cues, playful metaphors.
- Embed nostalgic or pop‑culture nods sparingly.
- Avoid duplicate phrasing across prompts.
- Use Asterisks around words e.g. ()"**signifant element**" OR capitalize certain words to highlight their significance (e.g. "act like GHOST")


***MANDATORY***
**ENSURE EACH PROMPT IS SEPERATED BY EXACTLY 3 HYPHENS E.G. "---"**
**ENSURE YOUR OUTPUT IS GENERATED IN A CODE BLOCK FOR EASY COPY AND PASTING. "---"**

Begin now, return only the Markdown prompts and delimiters.
"""


# ------------------------------------------------------------------ #
# Regex for bolding tokens with ≥2 uppercase letters
# ------------------------------------------------------------------ #
ACRONYM_RE = re.compile(r'.*[A-Z].*[A-Z].*')

def _add_paragraph_with_acronyms(p, text):
    """Split text into runs and bold tokens matching ACRONYM_RE."""
    for token in re.split(r'(\W+)', text):
        run = p.add_run(token)
        if ACRONYM_RE.fullmatch(token):
            run.bold = True

# ------------------------------------------------------------------ #
# Parse pasted Markdown into (title, description, body) tuples
# ------------------------------------------------------------------ #
def parse_prompts(raw_md: str):
    blocks = re.split(r'(?m)^\s*---\s*$', raw_md.strip())
    prompts = []
    for blk in blocks:
        blk = blk.strip()
        if not blk:
            continue
        lines = blk.splitlines()
        title = desc = ""
        body_lines = []
        # Extract UX TITLE and UX DESCRIPTION
        i = 0
        for i, ln in enumerate(lines):
            if ln.startswith("UX TITLE:"):
                title = ln.split("UX TITLE:", 1)[1].strip()
            elif ln.startswith("UX DESCRIPTION:"):
                desc = ln.split("UX DESCRIPTION:", 1)[1].strip()
                i += 1
                break
        # The rest is the body
        body_lines = lines[i:]
        prompts.append((title, desc, "\n".join(body_lines).strip()))
    return prompts

# ------------------------------------------------------------------ #
# Append prompts to a .docx with formatting and spacing
# ------------------------------------------------------------------ #
def append_to_docx(path: Path, prompts):
    try:
        doc = Document(path)
    except BadZipFile:
        if QMessageBox.question(
            None, "Invalid File",
            f"'{path.name}' is not a valid .docx.\nOverwrite with blank?"
        ) != QMessageBox.StandardButton.Yes:
            return False
        doc = Document()
        doc.save(path)
        doc = Document(path)

    for title, desc, body in prompts:
        # Page break before each prompt
        doc.add_page_break()

        # — Title —
        p_t = doc.add_paragraph()
        p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_t = p_t.add_run(title)
        run_t.font.name = 'Comic Sans MS'
        run_t.font.size = Pt(24)
        run_t.bold = True
        run_t.font.color.rgb = RGBColor(0x00, 0x80, 0x40)  # teal-green
        doc.add_paragraph()

        # — Description —
        p_d = doc.add_paragraph()
        run_d = p_d.add_run(desc)
        run_d.font.italic = True
        run_d.font.size = Pt(14)
        doc.add_paragraph()

        # — Copy instruction —
        p_c = doc.add_paragraph()
        p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_c = p_c.add_run("Copy entire prompt below:")
        run_c.font.bold = True
        run_c.font.size = Pt(11)
        doc.add_paragraph()

        # — Body —
        for ln in body.splitlines():
            s = ln.lstrip()
            # Page break for delimiter if standalone
            if s == '---':
                doc.add_page_break()
                continue

            # Numbered item → literal "1. …"
            m_num = re.match(r'^(\d+)\.\s+(.*)', s)
            if m_num:
                num, txt = m_num.groups()
                p = doc.add_paragraph(style='Normal')
                _add_paragraph_with_acronyms(p, f"{num}. {txt}")
                continue

            # Bullet item → Word bullet
            m_bul = re.match(r'^-\s+(.*)', s)
            if m_bul:
                txt = m_bul.group(1)
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.left_indent = Inches(0.25)
                _add_paragraph_with_acronyms(p, txt)
                continue

            # Plain paragraph
            p = doc.add_paragraph(style='Normal')
            _add_paragraph_with_acronyms(p, s)

        # Space after each prompt
        doc.add_paragraph()
        doc.add_paragraph()

    try:
        doc.save(path)
    except PermissionError:
        QMessageBox.critical(
            None, "Save Error",
            f"Cannot write to '{path.name}'. Close it and retry."
        )
        return False

    return True

# ------------------------------------------------------------------ #
# Main GUI application
# ------------------------------------------------------------------ #
class UnifiedApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("AI Prompt Toolkit")
        tabs = QTabWidget()
        tabs.addTab(self._master_tab(), "Master Prompt")
        tabs.addTab(self._append_tab(), "Append to Word")
        self.setCentralWidget(tabs)
        self.resize(700, 500)

    def _master_tab(self):
        w = QWidget(); layout = QVBoxLayout(w)
        # Chapter selector
        row = QHBoxLayout()
        row.addWidget(QLabel("Chapter:"))
        self.chapter_cb = QComboBox()
        self.chapter_cb.addItems(EXAMPLE_TITLES.keys())
        self.chapter_cb.setEditable(True)
        row.addWidget(self.chapter_cb)
        row.addStretch()
        layout.addLayout(row)
        # Prompt count selector
        row2 = QHBoxLayout()
        row2.addWidget(QLabel("Prompts:"))
        self.count_sb = QSpinBox()
        self.count_sb.setRange(1, 10)
        self.count_sb.setValue(5)
        row2.addWidget(self.count_sb)
        row2.addStretch()
        layout.addLayout(row2)
        # Copy button
        btn = QPushButton("Copy Master Prompt")
        btn.clicked.connect(self._copy_master)
        layout.addWidget(btn, alignment=Qt.AlignmentFlag.AlignLeft)
        layout.addStretch()
        return w

    def _append_tab(self):
        w = QWidget(); v = QVBoxLayout(w)
        # File selector
        fl = QHBoxLayout()
        self.file_label = QLabel("No document selected")
        select_btn = QPushButton("Select .docx…")
        select_btn.clicked.connect(self._select_doc)
        fl.addWidget(select_btn)
        fl.addWidget(self.file_label)
        fl.addStretch()
        v.addLayout(fl)
        # Open after save
        self.open_cb = QCheckBox("Open after save")
        self.open_cb.setChecked(False)
        v.addWidget(self.open_cb)
        # Paste area
        v.addWidget(QLabel("Paste Markdown prompts here:"))
        self.text_edit = QTextEdit()
        v.addWidget(self.text_edit)
        # Insert button
        ins_btn = QPushButton("Insert Prompts")
        ins_btn.clicked.connect(self._insert_prompts)
        v.addWidget(ins_btn, alignment=Qt.AlignmentFlag.AlignLeft)
        v.addStretch()
        return w

    def _copy_master(self):
        chap = self.chapter_cb.currentText().strip()
        cnt = self.count_sb.value()
        text = MASTER_TEMPLATE.format(
            chapter_name=chap,
            count=cnt
        )
        QGuiApplication.clipboard().setText(text)
        QMessageBox.information(self, "Copied",
            "Master prompt (Markdown) copied to clipboard."
        )

    def _select_doc(self):
        fn, _ = QFileDialog.getOpenFileName(
            self, "Select Word Document", "", "Word Documents (*.docx)"
        )
        if fn:
            self.doc_path = Path(fn)
            self.file_label.setText(self.doc_path.name)

    def _insert_prompts(self):
        if not hasattr(self, 'doc_path'):
            QMessageBox.warning(self, "No File", "Please select a .docx first.")
            return
        raw = self.text_edit.toPlainText().strip()
        if not raw:
            QMessageBox.warning(self, "Empty", "Nothing was pasted.")
            return

        prompts = parse_prompts(raw)
        if not prompts:
            QMessageBox.warning(self, "Parse Error",
                                "No valid prompts found. Ensure you have UX TITLE:, UX DESCRIPTION:, and --- separators.")
            return

        if append_to_docx(self.doc_path, prompts):
            if self.open_cb.isChecked():
                try: os.startfile(self.doc_path)
                except: pass
            self.text_edit.clear()
            QMessageBox.information(self, "Done", "Prompts appended successfully.")

# ------------------------------------------------------------------ #
# Optional dark theme
# ------------------------------------------------------------------ #
def apply_dark_theme(app):
    app.setStyle("Fusion")
    p = QPalette()
    p.setColor(QPalette.ColorRole.Window,       QColor(45,45,45))
    p.setColor(QPalette.ColorRole.WindowText,   QColor(220,220,220))
    p.setColor(QPalette.ColorRole.Base,         QColor(60,60,60))
    p.setColor(QPalette.ColorRole.AlternateBase,QColor(45,45,45))
    p.setColor(QPalette.ColorRole.ToolTipBase,  QColor(220,220,220))
    p.setColor(QPalette.ColorRole.ToolTipText,  QColor(220,220,220))
    p.setColor(QPalette.ColorRole.Text,         QColor(220,220,220))
    p.setColor(QPalette.ColorRole.Button,       QColor(60,60,60))
    p.setColor(QPalette.ColorRole.ButtonText,   QColor(220,220,220))
    p.setColor(QPalette.ColorRole.Highlight,    QColor(0,120,215))
    p.setColor(QPalette.ColorRole.HighlightedText, QColor(255,255,255))
    app.setPalette(p)

# ------------------------------------------------------------------ #
# Entry point
# ------------------------------------------------------------------ #
def main():
    app = QApplication(sys.argv)
    apply_dark_theme(app)
    win = UnifiedApp()
    win.show()
    sys.exit(app.exec())

if __name__ == "__main__":
    main()
