import sys
import os
import zipfile
import re
from pathlib import Path
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QTabWidget,
    QVBoxLayout, QHBoxLayout, QLabel, QComboBox,
    QSpinBox, QPushButton, QTextEdit, QCheckBox,
    QFileDialog, QMessageBox
)
from PyQt6.QtGui import QGuiApplication, QPalette, QColor
from PyQt6.QtCore import Qt
from zipfile import BadZipFile
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ------------------------------------------------------------------ #
# Example titles per chapter (unchanged)
# ------------------------------------------------------------------ #
EXAMPLE_TITLES = {
    "Quiz & Trivia Challenges": [
        "’80s Classroom Pop‑Quiz",
        "C‑3PO’s Galactic Trivia Circuit",
        "Saturday‑Morning Cartoon Quiz"
    ],
    "Immersive Role-Play Realms": [
        "Noir Detective Interview Roleplay",
        "Medieval Council Diplomacy Simulation",
        "Cyberpunk Undercover Negotiation"
    ],
    "Choose-Your-Path Adventures": [
        "Haunted Mansion Survival Branching Tale",
        "Solar Explorer’s Moral Crossroads",
        "Time‑Loop Detective Narrative"
    ],
    "Mystery & Logic Puzzles": [
        "Locked‑Room Alibi Reconstruction",
        "Ciphered Journal of a Mad Scientist",
        "Quantum Paradox Puzzle Chamber"
    ],
    "Storycraft Forge": [
        "Epic Myth Weaving Workshop",
        "Post‑Apocalypse Colony Chronicle",
        "Shakespearean Tragedy Remix"
    ],
    "What-If Simulators": [
        "If Dinosaurs Roamed Today Scenario",
        "Alternate History Space Race",
        "Steampunk Industrial Revolution Reimagined"
    ],
    "Language Playgrounds": [
        "Shakespearean Tongue Twister Challenge",
        "Emoji Pictionary Showdown",
        "Cryptic Crossword Duel"
    ],
    "Creative Brain Boosters": [
        "Surrealist Image Association Game",
        "Random Word Story Sprint",
        "One‑Sentence Microfiction Relay"
    ],
    "Mythic & Fantasy Escapes": [
        "Dragon‑Rider’s Riddle Tournament",
        "Faerie Market Negotiation Roleplay",
        "Wizard’s Duel of Arcane Trivia"
    ],
    "Social & Party Games": [
        "Speedy Movie Quote Face‑Off",
        "Musical Mashup Guessing Game",
        "Trivia Charades Relay"
    ],
}

# ------------------------------------------------------------------ #
# NEW master‑prompt template (system‑style)
# ------------------------------------------------------------------ #
MASTER_TEMPLATE = r"""IDENTITY & PURPOSE
You are an expert creative‑writing engine hired to supply **system prompts** that instruct a host LLM to run **extremely immersive experiences** other wordly experiences, hypoethetical scenarios, fun and interactive sessions, etc.  
For the chapter “{chapter_name},” create **{count}** completely unique system prompts, each 200‑300 words.

CORE DIRECTIVES (apply to every prompt)
- **Address the LLM, not the Player.** Speak in imperatives: “Adopt a pirate persona…”, “When the Player chooses an island, present …”. Never say “You (the Player)”.
- **Distinct VOCAL STYLES:** Every prompt assigns the LLM a fresh persona (pirate, glitchy AI, ’90s host, noir detective, cartoon character from Warner Brothers, character from Star Wars or the MCU, Historical Figure etc.).
- **Interactive MECHANICS:** Detail how the AI should randomise questions, track score/time, offer hints, branch outcomes, escalate stakes, remember choices, etc.
- **No Boilerplate:** Vary heading names, structure, length, tone—each prompt must feel handcrafted.
- **Inclusivity & PG‑13.**
- **ENSURE THE PROMPT INCLUDES INSTRUCTIONS FOR LMM TO ***NOT BREAK CHARACTER OR STRAY FROM THE EXPERIENCE*** REGARDLESS OF WHAT THE USER SAYS**
- **Create very vibrant and creative scenarios from ficitious to 

SECTIONS
**You do not need to include all section, taylor the prompt to the experience and make each prompt unique containing varying sections**
(Feel free to rename each section or include other sections e.g. INSTRUCTIONS, IDENTITY & ROLE, GOAL.)
1. **AI PERSONA / VOICE** – dialect, quirks, sound effects the LLM should adopt.  
2. **SESSION SETUP** – greeting script, initial choices, what to reveal or conceal.  
3. **QUESTION FLOW & ADAPTIVE LOGIC** – how to select/scale questions, manage timers, track state.  
4. **HINT / REWARD / PENALTY SYSTEMS** – currencies, limits, consequences.  
5. **END‑STATE CONDITIONS** – victory, partial success, failure, epilogue instructions.

**THESE SECTIONS SHOULD BE TAYLORED TO THE EXPERIENCE THAT THE PROMPT WILL CREATE.**

STYLE NOTES
- Active voice, present tense, vivid sensory cues, playful metaphors.
- Embed nostalgic or pop‑culture nods sparingly.
- Avoid duplicate phrasing across prompts.

OUTPUT FORMAT (for the host LLM)
- Output prompts in chat, separated by `=== NEXT PROMPT ===`.
- MANDATORY - Each title MUST be prefixed with "UX TITLE:"
- MANDATORY - Each description MUST be prefixed with "UX DESCRIPTION:"
- No code fences, no commentary—only the prompts and delimiters.
- EACH PROMPT SHOULD HAVE A RELEVANT TITLE WITH SIMPLE NAMING ALLOWING THE USER USER TO EASILY UNDERSTAND WHAT TO EXPECT BASED ON JUST THE TITLE OF EACH PROMPT.
- EACH PROMPT SHOULD BE PROPERLY FORMATTED WITH MARKDOWN AND BULLET POINTS INCLUDING "**" TO EMPHASIZE KEY ELEMENTS IN THE PROMPT BUT AVOID CODE BLOCKS.
- EACH PROMPT SHOULD BE FORMATTED IN THIS CHAT WINDOW SUCH THAT I CAN EASILY COPY AND PASTE EACH GENERATED PROMPTS TO USE IN ANOTHER LLM.
- MANDATORY - Use "-- " where appropriate (two hyphens followed by a space) to mark bullet points inside the prompt body. Example: -- This is a bullet item.
- MANDATORY - Use "--- " where appropriate (three hyphens followed by a space) to mark a numbered list inside the prompt body. Example: --- This is a numbered list.
- Use Asterisks around words and capitalize certain words to highlight their significance.
- Use both bullet points and numbered lists within the prompts to format them.
- Ensure you leave space and format the prompts in the above format but make them feel like a mark down format.
- DO NOT USE EMOJIS.

EXAMPLE TITLES (for inspiration only)
{example_list}
*(Use these as idea sparks; craft your own titles.)*
*INCLUDE THE TITLE NAME PREFIXED WITH *UX TITLE:*  AND BENEATH A SHORT DESCRIPTION PREFIXED WITH *UX DESCRIPTION* OF WHAT THE EXPERIENCE ENTAILS AND HOW TO USE IT"

Begin generating now; return only the prompts and delimiters."""

# ---------- PARSING UTILS -------------------------------------------- #
PROMPT_DELIM = "=== NEXT PROMPT ==="
END_DELIM    = "=== END OF PROMPTS ==="
TITLE_TAG    = "UX TITLE:"
DESC_TAG     = "UX DESCRIPTION:"
ACRONYM_RE   = re.compile(r'\b(?=\w*[A-Z].*[A-Z])\w+\b')

def parse_prompts(raw: str):
    blocks = [b.strip() for b in raw.split(PROMPT_DELIM)]
    for block in blocks:
        # remove any stray END_DELIM markers
        block = block.replace(END_DELIM, "").strip()
        if not block:
            continue
        lines = block.splitlines()
        title = desc = None
        start_idx = 0
        for i, ln in enumerate(lines):
            if ln.startswith(TITLE_TAG):
                title = ln.split(TITLE_TAG, 1)[1].strip()
            elif ln.startswith(DESC_TAG):
                desc = ln.split(DESC_TAG, 1)[1].strip()
                start_idx = i + 1
                break
        body = "\n".join(lines[start_idx:]).strip()
        yield (
            title or "(Untitled Prompt)",
            desc or "(No description provided)",
            body
        )

# ---------- WORD APPEND LOGIC ---------------------------------------- #
def append_to_docx(path: Path, prompts):
    try:
        doc = Document(path)
    except BadZipFile:
        if QMessageBox.question(
            None, "Invalid File",
            f"'{path.name}' is not a valid .docx.\nOverwrite with blank?"
        ) != QMessageBox.StandardButton.Yes:
            return False
        doc = Document()
        doc.save(path)
        doc = Document(path)

    # tighten Normal
    pf = doc.styles['Normal'].paragraph_format
    pf.space_before = Pt(0); pf.space_after = Pt(0); pf.line_spacing = 1
    # tighten headings
    for style in ('Heading 1','Heading 2'):
        pf = doc.styles[style].paragraph_format
        pf.space_before = Pt(0); pf.space_after = Pt(0); pf.line_spacing = 1

    for title, desc, body in prompts:
        doc.add_page_break()
        # Title
        pt = doc.add_paragraph()
        pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pt.add_run(title); run.bold=True; run.font.size=Pt(20)
        run.font.color.rgb = RGBColor(0x00,0x80,0x80)
        pt.style = 'Heading 1'
        doc.add_paragraph(); doc.add_paragraph()
        # Description
        pd = doc.add_paragraph()
        pd.alignment = WD_ALIGN_PARAGRAPH.LEFT
        rd = pd.add_run(desc); rd.bold=True; rd.italic=True
        pd.style = 'Normal'
        doc.add_paragraph()
        # Instruction
        pi = doc.add_paragraph()
        ri = pi.add_run("Copy the entire prompt below."); ri.font.size = Pt(10)
        pi.alignment = WD_ALIGN_PARAGRAPH.LEFT; pi.style = 'Normal'
        doc.add_paragraph(); doc.add_paragraph()
        # Body
        counter = 1
        for ln in body.splitlines():
            s = ln.lstrip()
            if s.startswith("--- "):
                txt = s[4:].strip()
                p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.25)
                p.add_run(f"{counter}. "); counter += 1
            elif s.startswith("-- "):
                txt = s[3:].strip()
                p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.25)
                p.add_run("• ")
            else:
                txt = ln; p = doc.add_paragraph()
            for tok in re.split(r'(\W+)', txt):
                r = p.add_run(tok)
                if ACRONYM_RE.fullmatch(tok): r.bold = True
        doc.add_paragraph(); doc.add_paragraph()

    try:
        doc.save(path)
    except PermissionError:
        QMessageBox.critical(
            None, "Save Error",
            f"Cannot write to '{path.name}'.\nClose it and retry."
        )
        return False
    return True

# ---------- UNIFIED GUI ------------------------------------------------ #
class UnifiedApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("AI Prompt Toolkit")
        tabs = QTabWidget()
        tabs.addTab(self._master_tab(), "Master Prompt")
        tabs.addTab(self._append_tab(), "Append to Word")
        self.setCentralWidget(tabs)
        self.resize(700, 500)

    def _master_tab(self):
        w = QWidget(); layout = QVBoxLayout(w)
        hl = QHBoxLayout()
        hl.addWidget(QLabel("Chapter:"))
        self.chapter_cb = QComboBox(); self.chapter_cb.addItems(EXAMPLE_TITLES.keys())
        self.chapter_cb.setEditable(True)
        hl.addWidget(self.chapter_cb); hl.addStretch()
        layout.addLayout(hl)
        hl2 = QHBoxLayout()
        hl2.addWidget(QLabel("Prompts:"))
        self.count_sb = QSpinBox(); self.count_sb.setRange(1,10); self.count_sb.setValue(5)
        hl2.addWidget(self.count_sb); hl2.addStretch()
        layout.addLayout(hl2)
        btn = QPushButton("Copy Master Prompt"); btn.clicked.connect(self._copy_master)
        layout.addWidget(btn, alignment=Qt.AlignmentFlag.AlignLeft)
        layout.addStretch()
        return w

    def _append_tab(self):
        w = QWidget(); v = QVBoxLayout(w)
        fl = QHBoxLayout()
        self.file_label = QLabel("No document selected")
        btn = QPushButton("Select .docx…"); btn.clicked.connect(self._select_doc)
        fl.addWidget(btn); fl.addWidget(self.file_label); fl.addStretch()
        v.addLayout(fl)
        self.open_cb = QCheckBox("Open after save"); self.open_cb.setChecked(False)
        v.addWidget(self.open_cb)
        v.addWidget(QLabel("Paste prompts:"))
        self.text_edit = QTextEdit(); v.addWidget(self.text_edit)
        btn2 = QPushButton("Insert Prompts"); btn2.clicked.connect(self._insert_prompts)
        v.addWidget(btn2, alignment=Qt.AlignmentFlag.AlignLeft)
        v.addStretch()
        return w

    def _copy_master(self):
        chap = self.chapter_cb.currentText().strip()
        cnt  = self.count_sb.value()
        ex   = EXAMPLE_TITLES.get(chap, [])
        el   = "\n".join(f"- {e}" for e in ex)
        text = MASTER_TEMPLATE.format(chapter_name=chap, count=cnt, example_list=el)
        QGuiApplication.clipboard().setText(text)

    def _select_doc(self):
        fn, _ = QFileDialog.getOpenFileName(self, "Select Word Document", "", "Word Documents (*.docx)")
        if fn:
            self.doc_path = Path(fn)
            self.file_label.setText(self.doc_path.name)

    def _insert_prompts(self):
        if not hasattr(self, 'doc_path'):
            QMessageBox.warning(self, "No File", "Please select a .docx first.")
            return
        raw = self.text_edit.toPlainText().strip()
        if not raw:
            QMessageBox.warning(self, "Empty", "Nothing was pasted.")
            return
        prompts = list(parse_prompts(raw))
        if not prompts:
            QMessageBox.warning(self, "Parse Error", "No valid prompts found.")
            return
        if not append_to_docx(self.doc_path, prompts):
            return
        if self.open_cb.isChecked():
            try: os.startfile(self.doc_path)
            except: pass
        # keep GUI open, clear for next batch
        self.text_edit.clear()

# ---------- DARK THEME ------------------------------------------------ #
def apply_dark_theme(app):
    app.setStyle("Fusion")
    p = QPalette()
    p.setColor(QPalette.ColorRole.Window,       QColor(45,45,45))
    p.setColor(QPalette.ColorRole.WindowText,   QColor(220,220,220))
    p.setColor(QPalette.ColorRole.Base,         QColor(60,60,60))
    p.setColor(QPalette.ColorRole.AlternateBase,QColor(45,45,45))
    p.setColor(QPalette.ColorRole.ToolTipBase,  QColor(220,220,220))
    p.setColor(QPalette.ColorRole.ToolTipText,  QColor(220,220,220))
    p.setColor(QPalette.ColorRole.Text,         QColor(220,220,220))
    p.setColor(QPalette.ColorRole.Button,       QColor(60,60,60))
    p.setColor(QPalette.ColorRole.ButtonText,   QColor(220,220,220))
    p.setColor(QPalette.ColorRole.Highlight,    QColor(0,120,215))
    p.setColor(QPalette.ColorRole.HighlightedText, QColor(255,255,255))
    app.setPalette(p)

def main():
    app = QApplication(sys.argv)
    apply_dark_theme(app)
    win = UnifiedApp()
    win.show()
    sys.exit(app.exec())

if __name__ == "__main__":
    main()
