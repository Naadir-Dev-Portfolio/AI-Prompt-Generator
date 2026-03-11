"""
Append pasted AI prompts to a Word document,
then automatically open it in Word when done.

Features:
- Centered, enlarged, colored titles
- Left-aligned italic descriptions
- Balanced “Copy the entire prompt below.” placement (slightly larger text)
- Automatic “-- ” bullets and “--- ” numbered bullets (incrementing)
- Acronym-style bolding
- Tightened line spacing
- Permission-error handling on save
- Launches the .docx after successful append

Workflow:
1. Select (or create) your .docx.
2. Paste prompts separated by '=== NEXT PROMPT ===',
   each containing 'UX TITLE:' and 'UX DESCRIPTION:'.
3. Click “Insert Prompts”: each gets its own formatted page,
   then the document opens automatically.

Dependencies:
    pip install python-docx
"""

import sys, os, zipfile, re
from pathlib import Path
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
from zipfile import BadZipFile
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------- CONSTANTS ------------------------------------------------ #
PROMPT_DELIM = "=== NEXT PROMPT ==="
TITLE_TAG    = "UX TITLE:"
DESC_TAG     = "UX DESCRIPTION:"
ACRONYM_RE   = re.compile(r'\b(?=\w*[A-Z].*[A-Z])\w+\b')
# Set to True if you want Word to open automatically after inserting prompts
OPEN_ON_SAVE = False


# ---------- PARSING -------------------------------------------------- #
def parse_prompts(raw: str):
    blocks = [b.strip() for b in raw.split(PROMPT_DELIM) if b.strip()]
    for block in blocks:
        lines = block.splitlines()
        title = desc = None
        body_start = 0
        for i, line in enumerate(lines):
            if line.startswith(TITLE_TAG):
                title = line.split(TITLE_TAG, 1)[1].strip()
            elif line.startswith(DESC_TAG):
                desc = line.split(DESC_TAG, 1)[1].strip()
                body_start = i + 1
                break
        yield (
            title or "(Untitled Prompt)",
            desc   or "(No description provided)",
            "\n".join(lines[body_start:]).strip()
        )

# ---------- WORD MANIPULATION --------------------------------------- #
def append_to_docx(doc_path: Path, prompts):
    """
    Open or create the .docx, append each prompt on a fresh page,
    apply formatting, then save. Handles permission errors.
    """
    try:
        doc = Document(doc_path)
    except BadZipFile:
        create_new = messagebox.askyesno(
            "Invalid Word File",
            f"'{doc_path.name}' is not a valid .docx.\n"
            "Overwrite it with a new blank document?"
        )
        if not create_new:
            return False
        doc = Document()
        doc.save(doc_path)
        doc = Document(doc_path)

    # Tighten Normal style
    normal_pf = doc.styles['Normal'].paragraph_format
    normal_pf.space_before = Pt(0)
    normal_pf.space_after  = Pt(0)
    normal_pf.line_spacing = 1

    # Tighten headings
    for style in ('Heading 1', 'Heading 2'):
        pf = doc.styles[style].paragraph_format
        pf.space_before = Pt(0)
        pf.space_after  = Pt(0)
        pf.line_spacing = 1

    # Append content
    for title, desc, body in prompts:
        doc.add_page_break()

        # Title
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rt = p_title.add_run(title)
        rt.bold = True
        rt.font.size = Pt(20)
        rt.font.color.rgb = RGBColor(0x00, 0x80, 0x80)
        p_title.style = 'Heading 1'

        # Space under title
        doc.add_paragraph()
        doc.add_paragraph()

        # Description
        p_desc = doc.add_paragraph()
        p_desc.alignment = WD_ALIGN_PARAGRAPH.LEFT
        rd = p_desc.add_run(desc)
        rd.bold = True; rd.italic = True
        p_desc.style = 'Normal'

        # Extra space below description
        doc.add_paragraph()

        # Instruction line (slightly larger)
        p_instr = doc.add_paragraph()
        pi = p_instr.add_run("Copy the entire prompt below.")
        pi.font.size = Pt(10)
        p_instr.style = 'Normal'
        p_instr.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Extra space before prompt
        doc.add_paragraph()
        doc.add_paragraph()

        # Body with bullets & numbering
        counter = 1
        for ln in body.split("\n"):
            s = ln.lstrip()
            if s.startswith("--- "):
                txt = s[4:].rstrip()
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.add_run(f"{counter}. ")
                counter += 1
            elif s.startswith("-- "):
                txt = s[3:].rstrip()
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.add_run("• ")
            else:
                txt = ln
                p = doc.add_paragraph()

            for tok in re.split(r'(\W+)', txt):
                run = p.add_run(tok)
                if ACRONYM_RE.fullmatch(tok):
                    run.bold = True

        # Extra space after prompt
        doc.add_paragraph()
        doc.add_paragraph()

    # Save with permission handling
    try:
        doc.save(doc_path)
    except PermissionError:
        messagebox.showerror(
            "Save Error",
            f"Cannot write to '{doc_path.name}'.\n"
            "Please close it in Word and try again."
        )
        return False

    return True

# ---------- GUI ------------------------------------------------------ #
class PromptInserter:
    """Handles file selection, prompt input, and appending logic."""
    def __init__(self):
        self.root = tk.Tk()
        self.root.withdraw()
        self.doc_path = self.ask_docx_file()
        if not self.doc_path:
            sys.exit()
        self.build_text_window()

    def ask_docx_file(self):
        while True:
            path = filedialog.askopenfilename(
                title="Select or create the Word document to append",
                filetypes=[("Word Document","*.docx")]
            )
            if not path:
                return None
            p = Path(path)
            if p.suffix.lower() != ".docx":
                messagebox.showerror("Wrong File Type","Please pick a .docx file.")
                continue
            return p

    def build_text_window(self):
        self.win = tk.Toplevel()
        self.win.title("Paste Prompts → Click 'Insert Prompts'")
        self.win.geometry("750x450")

        tk.Label(
            self.win,
            text="Paste your prompt batch below, then click 'Insert Prompts'.",
            font=("Segoe UI",10,"bold")
        ).pack(anchor="w",padx=10,pady=(8,4))

        bf = tk.Frame(self.win)
        bf.pack(side="bottom",fill="x",pady=8)
        tk.Button(bf,text="Insert Prompts",command=self.process,width=16
        ).pack(side="left",padx=6,ipadx=6,ipady=2)
        tk.Button(bf,text="Cancel",command=self.quit,width=10
        ).pack(side="left",padx=6,ipadx=6,ipady=2)

        self.text = scrolledtext.ScrolledText(
            self.win,wrap=tk.WORD,font=("Segoe UI",10)
        )
        self.text.pack(side="top",expand=True,fill="both",padx=10,pady=(0,4))
        self.root.mainloop()

    def process(self):
        raw = self.text.get("1.0", tk.END).strip()
        if not raw:
            messagebox.showerror("Empty", "Nothing was pasted.")
            return

        prompts = list(parse_prompts(raw))
        if not prompts:
            messagebox.showerror("Parse Error", "No valid prompts found.")
            return

        success = append_to_docx(self.doc_path, prompts)
        if not success:
            # append_to_docx already shows an error dialog on failure
            return

        if OPEN_ON_SAVE:
            try:
                os.startfile(self.doc_path)
            except Exception:
                pass

        # silently exit on success
        self.quit()


    def quit(self):
        self.win.destroy()
        self.root.quit()

if __name__=="__main__":
    PromptInserter()
